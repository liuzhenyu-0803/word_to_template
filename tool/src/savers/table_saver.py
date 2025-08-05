"""
表格保存模块 - 将HTML表格内容写回Word文档
"""
from callback.callback import callback_handler
import os
import re
from bs4 import BeautifulSoup
from docx.document import Document
from docx.table import Table
from global_define.constants import ATTR_ORIGINAL_CONTENT, TABLE_FILE_PREFIX, HTML_FILE_EXTENSION

from global_define import constants

def get_table_content_hash(table):
    """获取表格内容的哈希值，用于内容去重"""
    content = []
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            row_content.append(cell.text.strip())
        content.append('|'.join(row_content))
    return '\n'.join(content)

async def save_tables(doc: Document, html_content: str) -> None:
    """
    将HTML表格中的内容更新到Word文档对象中对应的表格。
    此实现通过构建逻辑网格来正确处理合并单元格，并正确处理嵌套表格。
    """
    try:
        # 收集所有表格，包括嵌套表格，使用与提取阶段一致的深度优先遍历
        all_tables = _collect_all_tables_dfs(doc)
        
        soup = BeautifulSoup(html_content, 'html.parser')
        html_tables = soup.find_all('table', recursive=True) # 获取所有HTML表格
        
        if not html_tables:
            await callback_handler.output_callback("HTML内容中未找到任何表格。")
            return

        for table_index, html_table in enumerate(html_tables):
            if table_index >= len(all_tables):
                await callback_handler.output_callback(f"警告: Word文档中没有与HTML表格 {table_index + 1} 对应的表格。")
                continue

            doc_table = all_tables[table_index]
            
            # 打印doc_table的尺寸
            doc_rows = len(doc_table.rows)
            doc_cols = max(len(row.cells) for row in doc_table.rows) if doc_rows > 0 else 0
            await callback_handler.output_callback(f"Word表格 {table_index + 1} 尺寸: 行={doc_rows}, 列={doc_cols}")

            html_grid = _build_grid_from_html(html_table)
            
            # 打印html_grid的尺寸
            html_grid_rows = len(html_grid)
            html_grid_cols = max(len(row) for row in html_grid) if html_grid_rows > 0 else 0
            await callback_handler.output_callback(f"HTML逻辑网格 {table_index + 1} 尺寸: 行={html_grid_rows}, 列={html_grid_cols}")

            # 打印html_grid的详细内容
            await callback_handler.output_callback(f"HTML逻辑网格 {table_index + 1} 详细内容:")
            for r, row in enumerate(html_grid):
                row_content = []
                for c, cell in enumerate(row):
                    if cell:
                        cell_id = cell.get('data-cell-id', 'N/A')
                        cell_text = cell.get_text(strip=True)
                        if len(cell_text) > 20: # 截断过长的文本
                            cell_text = cell_text[:17] + "..."
                        row_content.append(f"({r},{c}) ID:{cell_id} Text:'{cell_text}'")
                    else:
                        row_content.append(f"({r},{c}) None")
                await callback_handler.output_callback(" | ".join(row_content))

            for r, row in enumerate(html_grid):
                for c, cell in enumerate(row):
                    if cell and cell.has_attr(ATTR_ORIGINAL_CONTENT):
                        try:
                            cell_text = cell.get_text(strip=True)
                            cell_id = cell.get('data-cell-id', 'N/A')
                            
                            # 特别追踪"Infrared image"单元格
                            if "Infrared image" in cell_text:
                                await callback_handler.output_callback(f"=== 追踪 'Infrared image' 单元格 ===")
                                await callback_handler.output_callback(f"HTML表格索引: {table_index}")
                                await callback_handler.output_callback(f"HTML单元格ID: {cell_id}")
                                await callback_handler.output_callback(f"HTML单元格内容: '{cell_text}'")
                                await callback_handler.output_callback(f"HTML网格位置: ({r}, {c})")
                                await callback_handler.output_callback(f"Word表格索引: {table_index}")
                                await callback_handler.output_callback(f"Word表格坐标: ({r}, {c})")
                                
                                # 检查colspan和rowspan
                                colspan = cell.get('colspan', '1')
                                rowspan = cell.get('rowspan', '1')
                                await callback_handler.output_callback(f"HTML单元格合并: colspan={colspan}, rowspan={rowspan}")
                            
                            if cell.get_text(strip=True) == "{Inspection Date}":
                                # 如果单元格内容为Inspection Date，跳过
                                await callback_handler.output_callback(f"保存表格时出错: 单元格内容为Inspection Date")
                            
                            # 检查坐标是否超出Word表格范围
                            if r >= doc_rows or c >= doc_cols:
                                await callback_handler.output_callback(f"警告: 表格 {table_index + 1} 的坐标 ({r}, {c}) 超出Word表格实际范围。")
                                continue # 跳过当前单元格，避免IndexError

                            doc_cell = doc_table.cell(r, c)
                            
                            # 特别追踪"Infrared image"单元格的Word映射
                            if "Infrared image" in cell_text:
                                doc_cell_text = doc_cell.text.strip()
                                await callback_handler.output_callback(f"Word单元格原始内容: '{doc_cell_text}'")
                                await callback_handler.output_callback(f"即将更新为: '{cell_text}'")
                            # 为了保留样式，我们修改第一个run的文本并删除其余的runs
                            new_text = cell.get_text(strip=True)
                            # 假设我们只处理单元格的第一个段落
                            # 清空单元格的第一个段落，然后添加新的文本
                            # 这种方法会丢失原始run的样式，但可以避免错误
                            p = doc_cell.paragraphs[0]
                            # 移除所有现有的run
                            for run in p.runs:
                                p._element.remove(run._element)
                            # 添加新的run
                            p.add_run(new_text)

                        except IndexError:
                            await callback_handler.output_callback(f"警告: 表格 {table_index + 1} 的坐标 ({r}, {c}) 导致IndexError。")
                        except Exception as inner_e:
                            await callback_handler.output_callback(f"保存表格 {table_index + 1} 单元格 ({r}, {c}) 时发生未知错误: {inner_e}")
        doc.save(constants.DEFAULT_TEMPLATE_DOC_PATH)
    except Exception as e:
        await callback_handler.output_callback(f"保存表格时出错: {e}")

def _collect_all_tables_dfs(doc: Document) -> list[Table]:
    """
    使用深度优先遍历收集文档中的所有表格，包括嵌套表格。
    遍历顺序与BeautifulSoup的find_all('table', recursive=True)保持一致。
    
    参数:
        doc: python-docx的Document对象
        
    返回:
        list[Table]: 按深度优先顺序排列的所有表格列表
    """
    all_tables = []
    visited_table_ids = set() # 用于存储已访问的表格ID，避免重复
    content_hashes = set() # 存储表格内容哈希，用于去重

    def _dfs_collect_tables(element):
        """深度优先递归收集表格"""
        # 如果当前元素有tables属性，遍历其直接包含的表格
        if hasattr(element, 'tables'):
            for table in element.tables:
                table_id = id(table)
                content_hash = get_table_content_hash(table) # 获取表格内容哈希
                
                if table_id not in visited_table_ids and content_hash not in content_hashes:
                    all_tables.append(table)
                    visited_table_ids.add(table_id)
                    content_hashes.add(content_hash) # 添加内容哈希到集合
                    
                    # 然后深度优先遍历该表格的所有单元格，查找嵌套表格
                    for row in table.rows:
                        for cell in row.cells:
                            _dfs_collect_tables(cell)
    
    # 从文档开始深度优先遍历
    _dfs_collect_tables(doc)
    return all_tables

def _build_grid_from_html(html_table):
    """
    将HTML表格（可能包含colspan/rowspan）转换为一个二维列表（网格）。
    每个网格单元格包含HTML单元格元素。
    使用recursive=False避免嵌套表格干扰。
    """
    grid = []
    for r_idx, row in enumerate(html_table.find_all('tr', recursive=False)):  # 只找直接子行
        # 确保当前行在网格中存在
        if len(grid) <= r_idx:
            grid.append([])

        for cell in row.find_all(['td', 'th'], recursive=False):  # 只找直接子单元格
            # 移动到因之前单元格的rowspan而占用的、第一个可用的列位置
            c_idx = 0
            while len(grid[r_idx]) > c_idx and grid[r_idx][c_idx] is not None:
                c_idx += 1
            
            rowspan = int(cell.get('rowspan', 1))
            colspan = int(cell.get('colspan', 1))

            # 将单元格放置在网格中，并根据rowspan/colspan填充
            for i in range(rowspan):
                for j in range(colspan):
                    # 确保目标行存在
                    if len(grid) <= r_idx + i:
                        grid.append([])
                    # 确保目标行有足够的列
                    while len(grid[r_idx + i]) <= c_idx + j:
                        grid[r_idx + i].append(None)
                    grid[r_idx + i][c_idx + j] = cell
    return grid