import docx

def get_table_content_hash(table):
    """获取表格内容的哈希值，用于内容去重"""
    content = []
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            row_content.append(cell.text.strip())
        content.append('|'.join(row_content))
    return '\n'.join(content)

def extract_tables_from_cell(cell, tables_dict, content_hashes):
    """递归提取单元格中的嵌套表格"""
    for table in cell.tables:
        table_id = id(table)
        if table_id not in tables_dict:
            # 检查内容是否重复
            content_hash = get_table_content_hash(table)
            if content_hash not in content_hashes:
                tables_dict[table_id] = table
                content_hashes.add(content_hash)
                
                # 递归查找更深层的嵌套表格
                for row in table.rows:
                    for cell_in in row.cells:
                        extract_tables_from_cell(cell_in, tables_dict, content_hashes)

def extract_all_tables(doc_path):
    doc = docx.Document(doc_path)
    tables_dict = {}  # 使用字典存储，key为table的id，value为table对象
    content_hashes = set()  # 存储表格内容哈希，用于去重
    
    # 提取顶层表格
    for table in doc.tables:
        table_id = id(table)
        if table_id not in tables_dict:
            # 检查内容是否重复
            content_hash = get_table_content_hash(table)
            if content_hash not in content_hashes:
                tables_dict[table_id] = table
                content_hashes.add(content_hash)
        
        # 递归查找每个顶层表格中的嵌套表格
        for row in table.rows:
            for cell in row.cells:
                extract_tables_from_cell(cell, tables_dict, content_hashes)
    
    return list(tables_dict.values())

def print_table(table, idx):
    print(f"表格 {idx+1}:")
    for row in table.rows:
        row_text = [cell.text.replace('\n', ' ') for cell in row.cells]
        print('\t'.join(row_text))
    print('-' * 40)

def get_table_dimensions(table):
    """获取表格的行数和列数"""
    num_rows = len(table.rows)
    num_cols = 0
    if num_rows > 0:
        # 获取第一行的单元格数量作为列数，因为合并单元格可能导致每行单元格数量不同
        num_cols = len(table.rows[0].cells)
    return num_rows, num_cols

def check_merged_cells_in_row(row):
    """检查给定行中是否有合并单元格"""
    # python-docx 库中，判断单元格是否合并需要检查其内部的 XML 结构
    # 一个单元格是否合并，可以通过检查其 `_tc` (table cell) 元素的 `gridSpan` (水平合并)
    # 和 `vMerge` (垂直合并) 属性来判断。
    # `gridSpan` 属性表示单元格横跨的列数，大于1表示水平合并。
    # `vMerge` 属性表示垂直合并，如果存在且值为 'restart' 或 'continue' 则表示垂直合并。
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr() # 获取或添加 tcPr 元素，其中包含合并信息

        # 检查水平合并
        if hasattr(tcPr, 'gridSpan') and tcPr.gridSpan is not None and tcPr.gridSpan.val > 1:
            return True
        
        # 检查垂直合并
        if hasattr(tcPr, 'vMerge') and tcPr.vMerge is not None:
            # vMerge 属性的值可以是 'restart' (合并区域的起始单元格) 或 'continue' (合并区域的延续单元格)
            # 只要存在 vMerge 属性，就表示该单元格参与了垂直合并
            return True
    return False

def find_value_in_tables(tables, value_to_find):
    """在所有表格中查找指定值，并返回其表格索引、行索引和列索引"""
    found_locations = []
    for table_idx, table in enumerate(tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if value_to_find in cell.text:
                    found_locations.append({
                        "table_index": table_idx,
                        "row_index": row_idx,
                        "col_index": col_idx,
                        "cell_text": cell.text.strip()
                    })
    return found_locations

if __name__ == "__main__":
    docx_path = "报告1.docx"
    all_tables = extract_all_tables(docx_path)
    print(f"共提取到 {len(all_tables)} 个表格（含嵌套表格）")
    for i, table in enumerate(all_tables):
        print_table(table, i)

    # 查找指定值
    search_value = "0.97"
    locations = find_value_in_tables(all_tables, search_value)

    if locations:
        print(f"\n找到 '{search_value}' 的位置:")
        for loc in locations:
            print(f"  表格索引: {loc['table_index'] + 1}, 行索引: {loc['row_index'] + 1}, 列索引: {loc['col_index'] + 1}, 单元格内容: '{loc['cell_text']}'")
    else:
        print(f"\n未找到 '{search_value}'")

    # 获取表格 5 的行数和列数
    if len(all_tables) >= 5:
        table_5 = all_tables[4] # 表格索引从0开始，所以表格5是all_tables[4]
        rows, cols = get_table_dimensions(table_5)
        print(f"\n表格 5 的行数: {rows}, 列数: {cols}")

        # 检查表格 5 的第一行是否有合并单元格
        if rows > 0:
            first_row_table_5 = table_5.rows[0]
            has_merged_cells = check_merged_cells_in_row(first_row_table_5)
            print(f"表格 5 的第一行是否有合并单元格: {'是' if has_merged_cells else '否'}")

            if has_merged_cells:
                print("表格 5 第一行的合并单元格详细信息:")
                for col_idx, cell in enumerate(first_row_table_5.cells):
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    
                    grid_span_val = 1
                    if hasattr(tcPr, 'gridSpan') and tcPr.gridSpan is not None:
                        grid_span_val = tcPr.gridSpan.val
                    
                    v_merge_val = None
                    if hasattr(tcPr, 'vMerge') and tcPr.vMerge is not None:
                        v_merge_val = tcPr.vMerge.val if hasattr(tcPr.vMerge, 'val') else 'present' # 'present' for cases where val attribute is missing but vMerge exists

                    if grid_span_val > 1 or v_merge_val is not None:
                        print(f"  列索引: {col_idx + 1}, 单元格内容: '{cell.text.strip()}'")
                        if grid_span_val > 1:
                            print(f"    水平合并 (gridSpan): {grid_span_val} 列")
                        if v_merge_val is not None:
                            print(f"    垂直合并 (vMerge): {v_merge_val}")
        else:
            print("表格 5 没有行。")
    else:
        print("\n报告中没有表格 5。")